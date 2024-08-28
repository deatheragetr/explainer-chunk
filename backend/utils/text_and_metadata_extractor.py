import io
from typing import Dict, Any, Tuple
import PyPDF2
import ebooklib
from ebooklib import epub
import json
import markdown
import csv
import openpyxl
from bs4 import BeautifulSoup
import docx


async def extract_text_and_metadata(
    file_content: bytes, file_type: str
) -> Tuple[str, Dict[str, Any]]:
    content_io = io.BytesIO(file_content)

    extractors = {
        "application/pdf": extract_from_pdf,
        "application/epub+zip": extract_from_epub,
        "application/json": extract_from_json,
        "text/markdown": extract_from_markdown,
        "text/plain": extract_from_text,
        "text/csv": extract_from_csv,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": extract_from_excel,
        "text/html": extract_from_html,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_from_docx,
    }

    extractor = extractors.get(file_type, extract_from_unknown)
    return await extractor(content_io)


async def extract_from_pdf(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    reader = PyPDF2.PdfReader(content_io)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    metadata = reader.metadata
    return text, metadata


async def extract_from_epub(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    book = epub.read_epub(content_io)
    text = ""
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        text += item.get_content().decode("utf-8") + "\n"
    metadata = {
        "title": book.get_metadata("DC", "title"),
        "creator": book.get_metadata("DC", "creator"),
        "language": book.get_metadata("DC", "language"),
    }
    return text, metadata


async def extract_from_json(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    data = json.load(content_io)
    text = json.dumps(data, indent=2)
    metadata = {
        "key_count": (
            len(data)
            if isinstance(data, dict)
            else len(data) if isinstance(data, list) else 0
        )
    }
    return text, metadata


async def extract_from_markdown(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    md_content = content_io.read().decode("utf-8")
    html = markdown.markdown(md_content)
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text()
    metadata = {
        "headings": [
            h.text for h in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])
        ]
    }
    return text, metadata


async def extract_from_text(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    text = content_io.read().decode("utf-8")
    metadata = {"line_count": text.count("\n") + 1}
    return text, metadata


async def extract_from_csv(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    csv_reader = csv.reader(io.TextIOWrapper(content_io))
    rows = list(csv_reader)
    text = "\n".join([",".join(row) for row in rows])
    metadata = {"row_count": len(rows), "column_count": len(rows[0]) if rows else 0}
    return text, metadata


async def extract_from_excel(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    workbook = openpyxl.load_workbook(content_io)
    text = ""
    sheet_data = {}
    for sheet in workbook.sheetnames:
        ws = workbook[sheet]
        sheet_text = "\n".join(
            [",".join([str(cell.value) for cell in row]) for row in ws.iter_rows()]
        )
        text += f"Sheet: {sheet}\n{sheet_text}\n\n"
        sheet_data[sheet] = {"row_count": ws.max_row, "column_count": ws.max_column}
    metadata: Dict[str, Dict[str, Any]] = {"sheets": sheet_data}
    return text, metadata


async def extract_from_html(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    soup = BeautifulSoup(content_io, "html.parser")
    text = soup.get_text()
    metadata = {
        "title": soup.title.string if soup.title else None,
        "headings": [
            h.text for h in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])
        ],
        "links": len(soup.find_all("a")),
    }
    return text, metadata


async def extract_from_docx(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    doc = docx.Document(content_io)
    text = "\n".join([para.text for para in doc.paragraphs])
    metadata = {
        "paragraph_count": len(doc.paragraphs),
        "section_count": len(doc.sections),
    }
    return text, metadata


async def extract_from_unknown(content_io: io.BytesIO) -> Tuple[str, Dict[str, Any]]:
    text = content_io.read().decode("utf-8", errors="ignore")
    metadata = {"file_size": len(text)}
    return text, metadata
